import re
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import docx
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


class FileUploader:
    def __init__(self):
        self.uploaded_file = None
        self.resume_text = None
        self.job_posting_text = None
        self.vectorized_documents = None
        self.tfidf_matrix = None
        self.top_keywords = None
        
    def get_resume_file(self):
        
        # Title section and file upload instructions
        st.subheader('Upload Resume: ')
        st.markdown("""
                    **File Upload Instructions:**  
                    - Please upload a PDF, Word Document or Plain Text file. Please do not upload other file types in order to avoid an error reading your resume.
                    - Only one file can be uploaded.
                    - If you would like to upload a new resume for submission, please click the clear file button.
                    """)
        
        # check session states
        if ('uploaded_resume_file' in st.session_state) and ('exported_resume_text' in st.session_state):
            
            # create variables from session state
            resume_file, resume_txt = st.session_state['uploaded_resume_file'], st.session_state['exported_resume_text']
            
            # Update class attributes
            self.uploaded_file, self.resume_text = resume_file, resume_txt
            
            # write info and return
            st.info('Resume file has already been uploaded.')
            return self.uploaded_file, self.resume_text
        
        # File Upload Section
        uploaded_file = st.file_uploader(
            label='Upload File: ',
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=False,
            key='FileUploaderGetResumeFile'
        )
    
        # Try to extract the text from the file
        try:
            
            # If file exists, set the filename to lower
            if uploaded_file is not None:
                file_name = uploaded_file.name.lower()
            
            # Variable will be updated after the text is extracted from the uploaded file.
            resume_text = None
            
            # check file type for pdf
            if file_name.endswith('.pdf'):
                
                # Instantiate reader object
                reader = PdfReader(uploaded_file)
                
                # empty text variable to store our text
                text = ""
                
                # Iterate through each page of the pdf file and append the text to the text variable
                for page in reader.pages:
                    page_text = page.extract_text() # extract the text using the extract_text function
                    if page_text: # if it exists, append to text variable and add a new line after each addition
                        text += page_text + '\n'
                        
                # Update resume text
                resume_text = text
         
            # check file type for word doc
            elif file_name.endswith('.docx'):
                
                # Instantiate the Document object and pass in the uploaded file
                document = docx.Document(docx=uploaded_file)
                
                # Creat an empty list variable to append the extracted text to
                text_list = []
                
                # Iterate through each paragraph and append the text
                for paragraph in document.paragraphs:
                    text_list.append(paragraph.text) # Append the paragraph to the empty list
                    
                # Once we are done iterating, join each line paragraph in our text list by a newline
                final_text = "\n".join(text_list)
                
                # Update resume text
                resume_text = final_text
                
            # check if file type is a simple text document
            elif file_name.endswith('.txt'):
                text = uploaded_file.getvalue().decode('utf-8')
                resume_text = text
            
            # else invalid file type passed in. return an error
            else:
                # st.error('Invalid file type passed in. Please remove or clear the uploaded file, and try again.')
                return None
            
            # If resume_text has been updated and is not None, then update the class attribute, write a sample, create a session state and return the text
            if resume_text is not None:
                
                # write a sample of the text
                st.write(resume_text)
                st.success('Resume successfully uploaded!')
                
                # Update class attribute
                self.uploaded_file = uploaded_file
                self.resume_text = resume_text
                
                st.session_state['uploaded_resume_file'] = self.uploaded_file
                st.session_state['exported_resume_text'] = self.resume_text
                
                # return the resume text
                return self.uploaded_file, self.resume_text
                
        except Exception as e:
            # st.error(f'Unable to retrieve and read the resume. Received Error: {e}')
            return None

    def get_job_posting_file(self):
        
        st.subheader('Upload Job Posting File or Paste Text: ')
        st.markdown("""
                    **Job Posting Upload Instructions:**  
                    - You may either upload a PDF, Word Document or Text file containing the job posting.
                    - You may also just copy and paste the job posting in the text area below.
                    """)
        
        # check if session state exists
        if ('uploaded_job_posting' in st.session_state):
            job_posting_txt = st.session_state['uploaded_job_posting']
            self.job_posting_text = job_posting_txt
            st.write('\n'), st.write('\n'), st.write('\n'), st.write('\n'), st.write('\n')
            st.info('Job Posting file has already been uploaded.')
            return self.job_posting_text
        
        
        file_upload = st.columns(1)
        text_upload = st.columns(1)
        
        with file_upload[0]:
            uploaded_job_posting_file = st.file_uploader(
                label='Upload Job Posting: ',
                type=['pdf', 'docx', 'txt'],
                key='JobPostingFileUploader',
                accept_multiple_files=False
            )
        
        with text_upload[0]:
            text_job_posting = st.text_area(
                label='Paste Job Posting Here: ',
                value=' ',
                key='JobPostingTextAreaInput',
                placeholder='Input job posting here...'
            )
            
        # Check first if user uploaded text of the job posting; if text_job_posting exists
        if isinstance(text_job_posting, str) and len(text_job_posting) > 1000:
            # Update class attribute, create session state and return 
            self.job_posting_text = text_job_posting.strip().lower()
            
            st.info('Job Posting successfully uploaded.')
            st.session_state['uploaded_job_posting'] = self.job_posting_text
            return self.job_posting_text
        
        # Then check if user uploaded a file instead
        elif uploaded_job_posting_file:
            
            # Variabel that holds out text
            job_posting_text = None
            
            # create file name variable and set it to lowercase
            file_name = uploaded_job_posting_file.name.lower()
            
            # check if it ends with a pdf
            if file_name.endswith('.pdf'):
                
                # instantiate pdf reader object, and pass in the job posting file
                reader = PdfReader(stream=uploaded_job_posting_file)
                
                # empty text variable to store the text that will be pulled from the reader
                text = ""
                
                # Iterate through each page of the pdf file and append the text variable
                for page in reader.pages:
                    page_text = page.extract_text() # Pull the text out of the page
                    
                    # if page_text exists, append to the text variable
                    if page_text:
                        text += page_text + "\n"
                
                # Update the job posting text variable
                job_posting_text = text
            
            # check if it ends with a docx
            elif file_name.endswith('docx'):
                
                # instantiate the document object and pass in the uploaded file
                document = docx.Document(docx=uploaded_job_posting_file)
                
                # Create an empty list to store the text pulled from the document
                text = []
                
                # Iterate through each paragraph in the document and append it to the empty list
                for paragraph in document.paragraphs:
                    text.append(paragraph.text)
                    
                # After the list has been appended, join it all with a new line
                final_text = '\n'.join(text)
                
                # Update the job posting text variable
                job_posting_text = final_text
                
            # check if user uploaded a text file
            elif file_name.endswith('.txt'):
                text = uploaded_job_posting_file.getvalue().decode('utf-8')
                job_posting_text = text
                
            # else return none
            else:
                st.error('Invalid file type uploaded.')
                return None    
            
            # if the job posting text has successfully been updated
            if job_posting_text is not None:
                
                # Write a sample of the text
                st.write(job_posting_text)
                st.info('Job Posting file has successfully been uploaded.')
                
                # Strip all whitespaces and set to lower one more time
                job_posting_text = job_posting_text.strip().lower()
                
                # Update class attribute, create session state, and return
                self.job_posting_text = job_posting_text
                st.session_state['uploaded_job_posting'] = self.job_posting_text
                return self.job_posting_text
    
    def vectorize_resume_and_job_posting(self):
        
        # check both the resume and the job posting have been uploaded
        if (not self.resume_text) and (not self.job_posting_text):
            st.warning('Please upload your resume and the job posting to start the analyzer.')
            return None
        
        else:
            # make copies of the class attributes
            resume_copy, job_posting_copy = self.resume_text, self.job_posting_text
            
            # Vectorize the text
            try:
                st.subheader('Resume Analysis')
                
                vectorize_button = st.button(
                    label='Begin Resume Analysis',
                    key='ResumeJobPostingVectorizer'
                )
                
                # Once user clicks the button, vectorize the data
                if vectorize_button:
                    
                    # Create a TF-IDF vectorizer object and vectorize both objects
                    # The stop_words='english' parameter excludes common English words (like “the”, “and”, etc.).
                    vectorizer = TfidfVectorizer(stop_words='english')
                    documents = [resume_copy, job_posting_copy]
                    tfidf_matrix = vectorizer.fit_transform(documents)
                    
                    # Calculate the cosine similarity between the resume and the job posting
                    # tfidf_matrix[0:1] produces a matrix with just the first document, and comparing it against the full matrix returns an array of similarities. 
                    # [0][1] extracts the similarity value between the first and second document.
                    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix)[0][1]
                    
                    # Display the similarity
                    st.write(f"Resume Analysis Completed. Similarity Score: {similarity:.2f}")
                    st.markdown(f"""
                                **Resume Results:**
                                - **Similarity Score:** {similarity:.2f}
                                - A similarity score below 0.20 may require improvements to your resume. Please review the missing keywords, and update your resume accordingly.
                                - Excessive symbols, special characters and images in your resume can lead to poorer results. Keep it simple!
                                """)
                    
                    # Update class attribute 
                    self.vectorized_documents = vectorizer
                    self.tfidf_matrix = tfidf_matrix
                    
                    return self.vectorized_documents, self.tfidf_matrix
                    
            except Exception as e:
                # st.error(f'Unable to vectorize the resume and job posting. Received Error: {e}')          
                return None  
                
    def display_matches_and_keywords_between_resume_and_job_posting(self):
        
        if not self.vectorized_documents:
            st.write('')
            return None
        
        # copy the vectorized documents
        vectorizer_copy, tfidx_matrix_copy = self.vectorized_documents, self.tfidf_matrix
        
        # Identify the most important keywords
        try:
            
            # Extract the feature names
            feature_names = vectorizer_copy.get_feature_names_out()
            # st.write(feature_names)
            
            # Convert the job posting (second document) tfidf scores to a dense array
            job_posting_tfidf = tfidx_matrix_copy[1].toarray().flatten()
            # st.write(job_posting_tfidf)

            # Get the top 10 keywords, 
            N = 10
            top_indices = np.argsort(job_posting_tfidf)[-N:][::-1]
            top_keywords = [(feature_names[i], job_posting_tfidf[i]) for i in top_indices if job_posting_tfidf[i] > 0]
            
            # Update class attribute
            self.top_keywords = top_keywords
        
            # Display the keywords
            st.subheader('Most Important Keywords from Job Posting')
            keywords_df = pd.DataFrame(data=top_keywords, columns=['Keywords', 'Score'])
            keywords_df['Score (%)'] = (keywords_df['Score'] * 100).round(2)
            
            figure = px.bar(
                data_frame=keywords_df,
                x='Keywords',
                y='Score (%)',
                color='Score (%)',
                color_continuous_scale=px.colors.diverging.RdBu,
                title='Top 10 Keywords in the Job Posting',
                labels={
                    'Keywords': 'Important Words',
                    'Score (%)': 'Keyword Score in Percentage'
                }
            )
            
            st.plotly_chart(figure, use_container_width=True)
        
        except Exception as e:
            st.error(f'Unable to display the matches betwee the resume and job posting. Received Error: {e}')         
            return None
        
    def highlight_mising_keywords_in_resume(self):
        
        if not self.top_keywords:
            st.write('')
            return None
        
        # make a copy of the top keywords
        top_keywords = self.top_keywords.copy()
        
        # Find the missing keywords
        try:
            
            # Tokenize the resume
            resume_tokens = set(re.findall(r'\w+', self.resume_text.lower()))
            
            # Determine the missing keywords by checking which of the top keywords are missing from the resume tokens
            missing_keywords = [keyword for keyword, score in top_keywords if keyword.lower() not in resume_tokens]
            
            # Display the missing results to the user
            st.subheader('Missing Keywords from Resume')
            if missing_keywords:
                missing_df = pd.DataFrame(data=missing_keywords, columns=['Missing Keywords'])
                st.dataframe(missing_df, hide_index=True)
                
        except Exception as e:
            st.error(f"Unable to highlight the missing keywords. Received Error: {e}")   
            return None
        
        
        
        
        
        
        
        
        
        
        
        
        